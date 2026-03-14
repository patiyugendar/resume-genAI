import streamlit as st
import anthropic
import json
import io
import re
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib import colors

st.set_page_config(page_title="AI Resume Generator", page_icon="📄", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
.stApp{font-family:'Inter',sans-serif;background:#f4f6f9}
.title-block{background:linear-gradient(135deg,#1a1a2e,#16213e,#0f3460);padding:2rem 2.5rem;border-radius:16px;margin-bottom:2rem}
.title-block h1{color:white;font-size:2rem;margin:0}
.title-block p{color:#a0aec0;margin:.3rem 0 0;font-size:.95rem}
.section-card{background:white;border-radius:12px;padding:1.4rem;margin-bottom:1.1rem;box-shadow:0 2px 8px rgba(0,0,0,.06);border-left:4px solid #0f3460}
.resume-preview{background:white;border-radius:4px;padding:2.5rem 3rem;box-shadow:0 4px 20px rgba(0,0,0,.12);font-family:'Times New Roman','Times',serif;font-size:10pt;line-height:1.35;color:#000;max-width:860px;margin:0 auto}
.r-name{font-size:18pt;font-weight:700;text-align:center;color:#000;margin:0 0 4px 0;line-height:1.3;display:block}
.r-contact{text-align:center;font-size:9pt;color:#000;margin:0 0 5px 0;display:block}
.r-divider{border:none;border-top:1.5px solid #000;margin:4px 0 5px 0}
.r-section{font-size:12pt;font-weight:700;text-transform:uppercase;letter-spacing:.4px;color:#000;margin:7px 0 0 0;display:block}
.r-body{font-size:10pt;margin:1px 0;line-height:1.35;display:block}
.r-sub{font-size:10pt;font-weight:700;margin:4px 0 0 0;line-height:1.35;display:flex;justify-content:space-between;width:100%}
.r-italic{font-size:10pt;font-style:italic;margin:0 0 1px 0;line-height:1.35;display:block}
.r-bullet{font-size:10pt;margin:1px 0 1px 16px;line-height:1.35;display:block}
.keyword-badge{display:inline-block;background:#e8f4fd;color:#1a73e8;border:1px solid #1a73e8;border-radius:20px;padding:2px 10px;margin:3px;font-size:.8rem}
.match-badge{display:inline-block;background:#e6f4ea;color:#1e8e3e;border:1px solid #1e8e3e;border-radius:20px;padding:2px 10px;margin:3px;font-size:.8rem}
.miss-badge{display:inline-block;background:#fce8e6;color:#d93025;border:1px solid #d93025;border-radius:20px;padding:2px 10px;margin:3px;font-size:.8rem}
.warn-badge{display:inline-block;background:#fff3e0;color:#e65100;border:1px solid #e65100;border-radius:20px;padding:2px 10px;margin:3px;font-size:.8rem}
.stButton>button{background:linear-gradient(135deg,#0f3460,#1a73e8);color:white;border:none;border-radius:8px;padding:.6rem 2rem;font-size:1rem;font-weight:600;width:100%}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="title-block"><h1>📄 AI Resume Generator</h1><p>Powered by Claude · Job-Tailored · ATS-Optimized · Hallucination Detection · IE 5300 – Spring 2026</p></div>', unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
    st.caption("Your key is never stored.")
    st.markdown("---")
    st.markdown("### 📊 Prompt Version")
    prompt_version = st.selectbox("Prompt Strategy", ["v1 – Basic", "v2 – Structured", "v3 – Job-Tailored + ATS (Recommended)"], index=2)
    st.markdown("---")
    st.markdown("### 🛡️ Hallucination Detection")
    hallucination_check = st.toggle("Enable Hallucination Check", value=True)
    st.caption("Warns if AI adds info you didn't provide.")
    st.markdown("---")
    st.info("Fill in your details, paste a job description, and click Generate.")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 Build Resume", "📊 ATS Score Checker", "💬 Resume Chat", "📜 Prompt Log", "ℹ️ How to Use"])

# ═══════════════════════════════════════════
# SHARED HELPERS
# ═══════════════════════════════════════════
DATE_RE = re.compile(
    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*'
    r'[\s\-–]+\d{4}(?:\s*[-–]\s*(?:Present|[A-Z][a-z]+\s+\d{4}|\d{4}))?'
    r'|\d{4}\s*[-–]\s*(?:Present|\d{4}))',
    re.IGNORECASE
)

def clean_md(text):
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'^[-=]{2,}$', '', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&;', '', text)
    # Fix trailing semicolons after abbreviations: "A&M;" → "A&M", "AT&T;" → "AT&T"
    text = re.sub(r'([A-Z&]+);', lambda m: m.group(1), text)
    return text.strip()

def is_section_hdr(text):
    c = re.sub(r'[^A-Za-z\s]', '', text).strip()
    excluded = {'GPA','USA','PHD','MBA','BSC','MS','BS','NA','AWS','SEO','SEM','ROI','KPI'}
    return (c.isupper() and len(c) >= 4 and len(text) < 65
            and not text.startswith('•') and c.strip() not in excluded
            and not re.match(r'^[A-Z]{1,3}$', c.strip()))

def split_date(line):
    """Return (left_text, date_string) or (line, None)"""
    m = DATE_RE.search(line)
    if m:
        left = line[:m.start()].strip().rstrip('|,').strip()
        return left, m.group(0).strip()
    return line, None

def is_skill_cat(text):
    return ':' in text and not text.startswith('•') and len(text.split(':')[0]) < 45

# ═══════════════════════════════════════════
# HTML PREVIEW RENDERER
# ═══════════════════════════════════════════
def render_html(resume_text):
    lines = resume_text.split('\n')
    parts = ['<div class="resume-preview">']
    first = True; second = False; section = ""

    for raw in lines:
        ls = clean_md(raw)
        if not ls:
            parts.append('<div style="margin:2px 0"></div>'); continue

        if first:
            parts.append(f'<div class="r-name">{ls}</div>')
            first = False; second = True; continue

        if second:
            # Contact line — render at 9pt, centered
            parts.append(f'<div class="r-contact">{ls}</div>')
            parts.append('<hr class="r-divider">')
            second = False; continue

        if is_section_hdr(ls):
            section = ls.upper()
            parts.append(f'<div class="r-section">{ls}</div>')
            parts.append('<hr class="r-divider">')
            continue

        if ls.startswith('•') or ls.startswith('-'):
            b = ls.lstrip('-').strip()
            if not b.startswith('•'): b = '• ' + b
            parts.append(f'<div class="r-bullet">{b}</div>')
            continue

        if section in ("TECHNICAL SKILLS","SKILLS") and is_skill_cat(ls):
            ci = ls.index(':')
            cat = ls[:ci].strip(); rest = ls[ci+1:].strip()
            parts.append(f'<div class="r-bullet">• <span style="font-weight:700">{cat}:</span> {rest}</div>')
            continue

        if section == "EDUCATION":
            left, date = split_date(ls)
            if date:
                parts.append(f'<div class="r-sub"><span>{left}</span><span style="white-space:nowrap">{date}</span></div>')
            elif re.match(r'(?i)^gpa', ls):
                parts.append(f'<div class="r-italic">{ls}</div>')
            else:
                parts.append(f'<div class="r-body">{ls}</div>')
            continue

        if section in ("PROFESSIONAL EXPERIENCE","EXPERIENCE","WORK EXPERIENCE"):
            left, date = split_date(ls)
            if date:
                sub = [p.strip() for p in re.split(r'\s{2,}|\|', left) if p.strip()]
                if len(sub) >= 2:
                    title = sub[0]; company = ' '.join(sub[1:])
                    parts.append(f'<div class="r-sub"><span>{title}</span><span style="white-space:nowrap">{date}</span></div>')
                    parts.append(f'<div class="r-italic">{company}</div>')
                else:
                    parts.append(f'<div class="r-sub"><span>{left}</span><span style="white-space:nowrap">{date}</span></div>')
            else:
                parts.append(f'<div class="r-body">{ls}</div>')
            continue

        parts.append(f'<div class="r-body">{ls}</div>')

    parts.append('</div>')
    return '\n'.join(parts)

# ═══════════════════════════════════════════
# DOCX BUILDER — fixed tab stop & spacing
# ═══════════════════════════════════════════
def build_docx(resume_text):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    # Usable width = 8.5 - 0.75 - 0.75 = 7.0 inches
    USABLE = 6.8  # right-tab from left margin (7.0" text area, 6.8 leaves buffer for date text)

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(10)

    def sp(para, before=0, after=0):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)

    def add_border(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1');    bot.set(qn('w:color'),'000000')
        pBdr.append(bot); pPr.append(pBdr)

    def add_right_tab(para, inches):
        """Add a right-aligned tab stop at exact inch position."""
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        # Convert inches to twentieths of a point (twips): 1 inch = 1440 twips
        tab.set(qn('w:pos'), str(int(inches * 1440)))
        tabs.append(tab)
        pPr = para._p.get_or_add_pPr()
        pPr.append(tabs)

    def row_with_date(para, left_text, date_text, bold=True, italic_left=False):
        """Write title/degree on left, date right-aligned using tab stop."""
        sp(para, 3, 0)
        add_right_tab(para, USABLE)
        r1 = para.add_run(left_text)
        r1.bold = bold; r1.font.name = 'Times New Roman'; r1.font.size = Pt(10)
        if italic_left: r1.italic = True
        r2 = para.add_run('\t' + date_text)
        r2.bold = bold; r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)

    first = True; second = False; cur_section = ""

    for raw in resume_text.split('\n'):
        line = clean_md(raw)
        if not line:
            p = doc.add_paragraph(); sp(p, 0, 0); continue

        if first:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line); r.bold = True
            r.font.name = 'Times New Roman'; r.font.size = Pt(16)
            sp(p, 0, 2); first = False; second = True; continue

        if second:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line); r.font.name = 'Times New Roman'; r.font.size = Pt(9)
            sp(p, 0, 2); second = False; continue

        if is_section_hdr(line):
            cur_section = line.upper()
            p = doc.add_paragraph(); r = p.add_run(line)
            r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
            sp(p, 5, 1); add_border(p); continue

        if line.startswith('•') or line.startswith('-'):
            b = line.lstrip('-').strip()
            if not b.startswith('•'): b = '• ' + b
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            sp(p, 0, 0)
            r = p.add_run(b); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            continue

        if cur_section in ("TECHNICAL SKILLS","SKILLS") and is_skill_cat(line):
            ci = line.index(':')
            cat = line[:ci].strip(); rest = line[ci+1:].strip()
            p = doc.add_paragraph(); sp(p, 0, 0)
            p.paragraph_format.left_indent = Inches(0.2)
            rb = p.add_run(f'• {cat}:')
            rb.bold = True; rb.font.name = 'Times New Roman'; rb.font.size = Pt(10)
            rr = p.add_run(f' {rest}')
            rr.font.name = 'Times New Roman'; rr.font.size = Pt(10)
            continue

        left, date = split_date(line)
        if date:
            if cur_section in ("PROFESSIONAL EXPERIENCE","EXPERIENCE","WORK EXPERIENCE"):
                sub = [s.strip() for s in re.split(r'\s{2,}|\|', left) if s.strip()]
                title   = sub[0] if sub else left
                company = ' '.join(sub[1:]) if len(sub) > 1 else ''
                # Title + date on same line
                p = doc.add_paragraph()
                row_with_date(p, title, date, bold=True)
                # Company italic on next line
                if company:
                    pc = doc.add_paragraph(); sp(pc, 0, 0)
                    rc = pc.add_run(company)
                    rc.italic = True; rc.font.name = 'Times New Roman'; rc.font.size = Pt(10)
            elif cur_section == "EDUCATION":
                # Split: "M.S. Data Science   University of Texas" → degree | university
                sub = [s.strip() for s in re.split(r'\s{3,}', left) if s.strip()]
                if len(sub) >= 2:
                    degree = sub[0]
                    university = ' '.join(sub[1:])
                    p = doc.add_paragraph(); sp(p, 3, 0)
                    add_right_tab(p, USABLE)
                    rb = p.add_run(degree)
                    rb.bold = True; rb.font.name = 'Times New Roman'; rb.font.size = Pt(10)
                    rm = p.add_run(f'   {university}')
                    rm.bold = False; rm.font.name = 'Times New Roman'; rm.font.size = Pt(10)
                    rd = p.add_run(f'\t{date}')
                    rd.bold = True; rd.font.name = 'Times New Roman'; rd.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    row_with_date(p, left, date, bold=True)
            else:
                p = doc.add_paragraph()
                row_with_date(p, left, date, bold=True)
            continue

        if re.match(r'(?i)^gpa', line):
            p = doc.add_paragraph(); sp(p, 0, 0)
            r = p.add_run(line); r.italic = True
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            continue

        # Plain body lines (coursework, summary text, etc) — no indent, normal weight
        p = doc.add_paragraph(); sp(p, 0, 0)
        r = p.add_run(line)
        r.bold = False; r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ═══════════════════════════════════════════
# PDF BUILDER — fixed contact line & date alignment
# ═══════════════════════════════════════════
def build_pdf(resume_text):
    buf = io.BytesIO()
    # Page: 8.5 wide, margins 0.75 each side → usable = 7.0 inches
    PAGE_W = 7.0 * inch  # 8.5 - 0.75 left - 0.75 right = 7.0 usable

    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch)

    nameS = ParagraphStyle('nm',  fontName='Times-Bold',   fontSize=16, alignment=1,
                           spaceAfter=6,  spaceBefore=0,  leading=22)
    contS = ParagraphStyle('ct',  fontName='Times-Roman',  fontSize=9,  alignment=1,
                           spaceAfter=6,  spaceBefore=2,  leading=12)
    # Section header: no space after — HR line sits immediately below
    secS  = ParagraphStyle('sc',  fontName='Times-Bold',   fontSize=12, alignment=0,
                           spaceAfter=0,  spaceBefore=8,  leading=15)
    bodyS = ParagraphStyle('bd',  fontName='Times-Roman',  fontSize=10,
                           spaceAfter=1,  spaceBefore=0,  leading=13)
    # Sub heading (job title / degree): no spaceBefore — sits immediately under HR line
    subS  = ParagraphStyle('sb',  fontName='Times-Bold',   fontSize=10,
                           spaceAfter=0,  spaceBefore=0,  leading=13)
    itaS  = ParagraphStyle('it',  fontName='Times-Italic', fontSize=10,
                           spaceAfter=1,  spaceBefore=0,  leading=13)
    bulS  = ParagraphStyle('bl',  fontName='Times-Roman',  fontSize=10,
                           leftIndent=14, spaceAfter=1,   spaceBefore=0, leading=13)

    def make_date_row(left_text, date_text, style, space_before=4):
        """Two-column table: left text | right-aligned date. Fixed to PAGE_W."""
        from reportlab.platypus import KeepTogether
        t = Table(
            [[Paragraph(left_text, style), Paragraph(date_text, style)]],
            colWidths=[PAGE_W - 1.6*inch, 1.6*inch]  # date col fixed at 1.6in, never wraps
        )
        t.setStyle(TableStyle([
            ('ALIGN',         (1,0),(1,0), 'RIGHT'),
            ('VALIGN',        (0,0),(-1,-1),'TOP'),
            ('LEFTPADDING',   (0,0),(-1,-1), 0),
            ('RIGHTPADDING',  (0,0),(-1,-1), 0),
            ('TOPPADDING',    (0,0),(-1,-1), 0),
            ('BOTTOMPADDING', (0,0),(-1,-1), 0),
        ]))
        # Add spacer before each job/education entry (except first which follows HR)
        if space_before > 0:
            return [Spacer(1, space_before), t]
        return [t]

    story = []
    first = True; second = False; cur_section = ""

    for raw in resume_text.split('\n'):
        line = clean_md(raw)
        if not line:
            story.append(Spacer(1, 2)); continue

        if first:
            story.append(Paragraph(line, nameS))
            story.append(HRFlowable(width="100%", thickness=1.5,
                                    color=colors.black, spaceAfter=4, spaceBefore=0))
            first = False; second = True; continue

        if second:
            # Contact at 9pt — fits on one line
            story.append(Paragraph(line, contS))
            second = False; continue

        if is_section_hdr(line):
            cur_section = line.upper()
            story.append(Paragraph(line, secS))
            story.append(HRFlowable(width="100%", thickness=0.8,
                                    color=colors.black, spaceAfter=2, spaceBefore=1))
            continue

        if line.startswith('•') or line.startswith('-'):
            b = line.lstrip('-').strip()
            if not b.startswith('•'): b = '• ' + b
            story.append(Paragraph(b, bulS)); continue

        if cur_section in ("TECHNICAL SKILLS","SKILLS") and is_skill_cat(line):
            ci = line.index(':')
            cat = line[:ci].strip(); rest = line[ci+1:].strip()
            story.append(Paragraph(f'• <b>{cat}:</b> {rest}', bulS)); continue

        left, date = split_date(line)
        if date:
            if cur_section in ("PROFESSIONAL EXPERIENCE","EXPERIENCE","WORK EXPERIENCE"):
                sub = [s.strip() for s in re.split(r'\s{2,}|\|', left) if s.strip()]
                title   = sub[0] if sub else left
                company = ' '.join(sub[1:]) if len(sub) > 1 else ''
                # Check if previous item was an HR line (first entry) or a bullet (subsequent)
                last = story[-1] if story else None
                sb = 0 if isinstance(last, HRFlowable) else 6
                story.extend(make_date_row(title, date, subS, space_before=sb))
                if company:
                    story.append(Paragraph(company, itaS))
            else:
                last = story[-1] if story else None
                sb = 0 if isinstance(last, HRFlowable) else 4
                story.extend(make_date_row(left, date, subS, space_before=sb))
            continue

        if re.match(r'(?i)^gpa', line):
            story.append(Paragraph(line, itaS)); continue

        # Coursework and other plain lines — normal weight
        story.append(Paragraph(line, bodyS))

    doc.build(story)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════
# TAB 1 — Build Resume
# ═══════════════════════════════════════════
with tab1:
    col_left, col_right = st.columns([1, 1], gap="large")
    with col_left:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 👤 Personal Information")
        full_name = st.text_input("Full Name", value="")
        email     = st.text_input("Email", value="")
        phone     = st.text_input("Phone", value="")
        linkedin  = st.text_input("LinkedIn URL", value="")
        location  = st.text_input("Location", value="")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🎓 Education")
        education = st.text_area("Education", height=130, placeholder="e.g. M.S. Computer Science, UT Dallas, Aug 2021 - May 2023, GPA: 3.9\nCoursework: Machine Learning, Distributed Systems\n\nB.S. Engineering, Texas A&M, Aug 2017 - May 2021")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🛠️ Skills")
        skills = st.text_area("Skills", height=120, placeholder="e.g. Software & Tools: Python, SQL, Excel, Power BI\nAnalytical: Six Sigma, DMAIC, SPC\nCertifications: PMP, AWS Certified")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 💼 Work Experience")
        experience = st.text_area("Experience", height=180, placeholder="e.g. Software Engineer, Google, Jun 2022 - Present:\n- Built REST APIs serving 500K daily requests\n- Reduced latency by 40% using Redis caching\n- Led team of 4 engineers")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🚀 Projects")
        projects = st.text_area("Projects", height=120, placeholder="e.g. Customer Segmentation Model:\n- Clustered 500K users using K-Means, increasing ROI by 30%\n- Deployed on AWS serving real-time predictions")
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🏆 Certifications")
        certifications = st.text_area("Certifications", height=80, placeholder="e.g. AWS Certified Developer\nLean Six Sigma Green Belt\nPMP Certification")
        st.markdown('</div>', unsafe_allow_html=True)

    with col_right:
        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🎯 Target Job")
        target_title = st.text_input("Target Job Title", placeholder="e.g. Quality Engineer")
        job_desc = st.text_area("Paste Full Job Description Here",
            placeholder="Copy and paste the complete job posting here...", height=340)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-card">', unsafe_allow_html=True)
        st.markdown("#### 🎨 Style Options")
        tone = st.radio("Tone",
            ["Professional & Conservative", "Modern & Dynamic", "Technical & Detailed"],
            horizontal=True, index=2)
        resume_length = st.select_slider("Length",
            options=["Concise (1 page)", "Standard (1–1.5 pages)", "Detailed (2 pages)"],
            value="Standard (1–1.5 pages)")
        st.markdown('</div>', unsafe_allow_html=True)

        generate = st.button("✨ Generate Tailored Resume", use_container_width=True)

    if generate:
        if not api_key:
            st.error("⚠️ Please enter your Anthropic API key in the sidebar.")
        elif not full_name or not experience or not job_desc:
            st.warning("⚠️ Please fill in at least Name, Experience, and Job Description.")
        else:
            if "v1" in prompt_version:
                sys_p = "You are a resume writer. Write a professional resume."
                usr_p = f"Write a resume for {full_name}. Skills: {skills}. Experience: {experience}."

            elif "v2" in prompt_version:
                sys_p = "You are an expert resume writer. Create structured professional resumes with clear sections."
                usr_p = (f"Name: {full_name}\nEmail: {email}\nEducation: {education}\n"
                         f"Skills: {skills}\nExperience: {experience}\nProjects: {projects}\n"
                         f"Certifications: {certifications}\nTarget: {target_title}")
            else:
                sys_p = """You are a senior resume strategist and ATS expert.

CRITICAL FORMATTING RULES — FOLLOW EXACTLY:
1. Zero markdown: no ##, no **, no *, no --, no hashtags, no asterisks anywhere
2. Section headers in ALL CAPS plain text only
3. Bullets use • only
4. Experience lines: ALWAYS write Job Title and Company Name and Date ALL ON ONE SINGLE LINE like this:
   Job Title   Company Name   Mon YYYY - Mon YYYY
5. Education lines: ALWAYS write Degree and University and Date ALL ON ONE SINGLE LINE like this:
   Degree Name   University Name   Mon YYYY - Mon YYYY
6. NEVER put company name on its own separate line
7. If employment gap exists, note [Career Gap: Mon YYYY – Mon YYYY] beside the job title on same line

KEYWORD RULES:
- Extract the top 12-15 keywords and required skills from the job description
- Naturally weave ALL of these keywords into the experience and project bullet points
- Mirror the exact phrasing from the job description in the bullets
- Every bullet must include at least one keyword from the job description

OUTPUT STRUCTURE:
FULL NAME
City | Phone | Email | LinkedIn

PROFESSIONAL SUMMARY
[3 sentences using job description keywords]

EDUCATION
Degree Name   University Name   Start - End
GPA: X.X (if provided)
Coursework: subject1, subject2, subject3

TECHNICAL SKILLS
Category: skill1, skill2, skill3

PROFESSIONAL EXPERIENCE
Job Title   Company   Start - End
• keyword-rich bullet with metric
• keyword-rich bullet with metric

PROJECTS
Project Title
• keyword-rich bullet with metric

CERTIFICATIONS
• Certification name

ONLY use info provided. NEVER invent anything."""

                usr_p = (
                    f"Name: {full_name} | Email: {email} | Phone: {phone} | "
                    f"Location: {location} | LinkedIn: {linkedin}\n"
                    f"EDUCATION: {education}\n"
                    f"SKILLS: {skills}\n"
                    f"EXPERIENCE: {experience}\n"
                    f"PROJECTS: {projects}\n"
                    f"TARGET ROLE: {target_title} | TONE: {tone} | LENGTH: {resume_length}\n"
                    f"JOB DESCRIPTION:\n{job_desc}\n\n"
                    f"Generate resume now. Plain text only. No markdown whatsoever."
                )

            with st.spinner("🤖 Claude is crafting your resume..."):
                try:
                    client = anthropic.Anthropic(api_key=api_key)
                    resp = client.messages.create(
                        model="claude-opus-4-5", max_tokens=2500,
                        system=sys_p,
                        messages=[{"role": "user", "content": usr_p}]
                    )
                    resume_text = resp.content[0].text
                    st.session_state["resume_text"]     = resume_text
                    st.session_state["candidate_name"]  = full_name
                    st.session_state["job_title"]       = target_title
                    st.session_state.pop("hal_data", None)

                    if hallucination_check:
                        hp = (f"Hallucination detector.\n"
                              f"ORIGINAL: {skills} {experience} {projects} {education}\n"
                              f"RESUME: {resume_text}\n"
                              f'Return ONLY JSON: {{"hallucinations_found":false,"suspicious_items":[],"verdict":"one sentence"}}')
                        hr = client.messages.create(model="claude-opus-4-5", max_tokens=500,
                            messages=[{"role":"user","content":hp}])
                        raw = hr.content[0].text.strip()
                        if raw.startswith("```"):
                            raw = raw.split("```")[1]
                            if raw.startswith("json"): raw = raw[4:]
                        st.session_state["hal_data"] = json.loads(raw)
                except Exception as e:
                    st.error(f"API Error: {e}")

    if "resume_text" in st.session_state:
        resume_text = st.session_state["resume_text"]
        st.markdown("---")

        if "hal_data" in st.session_state:
            hal = st.session_state["hal_data"]
            if hal.get("hallucinations_found"):
                st.warning(f"⚠️ **Hallucination Warning:** {hal.get('verdict','')}")
                items = hal.get("suspicious_items",[])
                if items:
                    badges = " ".join([f'<span class="warn-badge">⚠️ {i}</span>' for i in items])
                    st.markdown(f"**Suspicious items:** {badges}", unsafe_allow_html=True)
                st.caption("Review carefully before using.")
            else:
                st.success(f"✅ **Hallucination Check Passed** — {hal.get('verdict','No invented content detected.')}")

        st.markdown("### 📄 Your Tailored Resume")
        st.markdown(render_html(resume_text), unsafe_allow_html=True)

        st.markdown("#### 📋 Copy Resume Text")
        st.text_area("Click inside → Ctrl+A → Ctrl+C", value=resume_text, height=100, key="copy_area")

        st.markdown("#### 📥 Download")
        cname = st.session_state.get("candidate_name","Resume")
        jtitle = st.session_state.get("job_title","Role")
        fname = f"{cname.replace(' ','_')}_{jtitle.replace(' ','_')}"

        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("📝 Download Word (.docx)",
                data=build_docx(resume_text), file_name=f"{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        with dl2:
            st.download_button("📕 Download PDF",
                data=build_pdf(resume_text), file_name=f"{fname}.pdf",
                mime="application/pdf", use_container_width=True)

# ═══════════════════════════════════════════
# TAB 2 — ATS Score Checker
# ═══════════════════════════════════════════
with tab2:
    st.markdown("### 📊 ATS Score Checker")
    if not api_key:
        st.info("⚠️ Enter your Anthropic API key in the sidebar.")
    else:
        c1, c2 = st.columns(2)
        with c1:
            ats_job = st.text_area("📋 Paste Job Description", height=250, key="ats_job",
                placeholder="Paste the full job description here...")
        with c2:
            ats_resume = st.text_area("📄 Paste Your Resume Text", height=250, key="ats_resume",
                value=st.session_state.get("resume_text",""),
                placeholder="Auto-filled from Tab 1, or paste manually...")

        if st.button("🔎 Check ATS Score", use_container_width=True):
            if ats_job and ats_resume:
                with st.spinner("Analyzing ATS compatibility..."):
                    try:
                        client = anthropic.Anthropic(api_key=api_key)
                        ap = f"""You are an expert ATS analyst. Carefully compare this resume to the job description.

JOB DESCRIPTION:
{ats_job}

RESUME:
{ats_resume}

Instructions:
- Extract ALL important keywords, tools, skills, qualifications from the job description
- Check each one individually against the resume text
- Be thorough — check for synonyms and partial matches too
- Return ONLY valid JSON, nothing else:

{{"ats_score": 85,
"required_keywords": ["keyword1","keyword2","keyword3","keyword4","keyword5","keyword6","keyword7","keyword8"],
"matched_keywords": ["keyword1","keyword2","keyword3","keyword4","keyword5"],
"missing_keywords": ["keyword6","keyword7","keyword8"],
"strengths": ["strength1","strength2","strength3"],
"improvements": ["improvement1","improvement2","improvement3"],
"summary": "Two sentence overall assessment."}}"""

                        r = client.messages.create(model="claude-opus-4-5", max_tokens=1200,
                            messages=[{"role":"user","content":ap}])
                        raw = r.content[0].text.strip()
                        if raw.startswith("```"):
                            raw = raw.split("```")[1]
                            if raw.startswith("json"): raw = raw[4:]
                        d = json.loads(raw)
                        score = d.get("ats_score", 0)
                        clr = "#1e8e3e" if score>=75 else "#f9ab00" if score>=50 else "#d93025"
                        lbl = "Strong Match ✅" if score>=75 else "Moderate Match ⚠️" if score>=50 else "Weak Match ❌"
                        st.markdown(
                            f'<div style="text-align:center;padding:1.8rem;background:white;border-radius:14px;'
                            f'box-shadow:0 4px 16px rgba(0,0,0,.1);margin:1rem 0">'
                            f'<div style="font-size:3.5rem;font-weight:800;color:{clr}">{score}%</div>'
                            f'<div style="color:{clr};font-size:1.1rem;font-weight:600">{lbl}</div>'
                            f'<div style="color:#666;font-size:.9rem">ATS Compatibility Score</div></div>',
                            unsafe_allow_html=True)
                        k1,k2,k3 = st.columns(3)
                        with k1:
                            st.markdown("**📌 Required Keywords**")
                            st.markdown(" ".join([f'<span class="keyword-badge">{k}</span>'
                                for k in d.get("required_keywords",[])]) or "None", unsafe_allow_html=True)
                        with k2:
                            st.markdown("**✅ Matched**")
                            st.markdown(" ".join([f'<span class="match-badge">{k}</span>'
                                for k in d.get("matched_keywords",[])]) or "None", unsafe_allow_html=True)
                        with k3:
                            st.markdown("**❌ Missing**")
                            st.markdown(" ".join([f'<span class="miss-badge">{k}</span>'
                                for k in d.get("missing_keywords",[])]) or "None", unsafe_allow_html=True)
                        st.markdown("---")
                        s1,s2 = st.columns(2)
                        with s1:
                            st.markdown("**💪 Strengths**")
                            for s in d.get("strengths",[]): st.markdown(f"✅ {s}")
                        with s2:
                            st.markdown("**🔧 Improvements**")
                            for i in d.get("improvements",[]): st.markdown(f"➡️ {i}")
                        st.info(f"📝 **Summary:** {d.get('summary','')}")

                        # Store ATS data in session
                        st.session_state["ats_data"] = d
                        st.session_state["ats_job_desc"] = ats_job
                        st.session_state["ats_resume_text"] = ats_resume

                    except Exception as e:
                        st.error(f"Error: {e}")
            else:
                st.warning("Please fill in both fields.")

# ═══════════════════════════════════════════
# TAB 3 — Resume Chat Assistant
# ═══════════════════════════════════════════
with tab3:
    st.markdown("### 💬 Resume Chat Assistant")
    st.markdown("Chat with Claude to refine your resume. Ask it to change tone, add keywords, shorten sections, and more.")

    if not api_key:
        st.info("⚠️ Enter your Anthropic API key in the sidebar.")
    elif "resume_text" not in st.session_state:
        st.warning("⚠️ Generate a resume in Tab 1 first — it will auto-load here.")
    else:
        # Show current resume being edited
        current_resume = st.session_state.get("chat_resume", st.session_state.get("resume_text", ""))

        # Quick action buttons
        st.markdown("#### ⚡ Quick Actions")
        qa_cols = st.columns(4)
        quick_actions = [
            ("💪 Stronger Verbs",    "Rewrite all bullet points with stronger, more impactful action verbs"),
            ("✂️ Make Concise",      "Shorten the resume to fit neatly on one page without losing key information"),
            ("🎯 More Keywords",     "Add more relevant industry keywords throughout to improve ATS score"),
            ("📢 Bolder Summary",    "Rewrite the professional summary to be more confident and compelling"),
        ]
        for col, (label, action) in zip(qa_cols, quick_actions):
            with col:
                if st.button(label, use_container_width=True, key=f"qa_{label}"):
                    st.session_state["chat_input_prefill"] = action

        st.markdown("---")

        # Chat history display
        if "chat_history" not in st.session_state:
            st.session_state["chat_history"] = []

        # Display chat messages
        for msg in st.session_state["chat_history"]:
            if msg["role"] == "user":
                st.markdown(
                    f'<div style="background:#e8f4fd;border-radius:12px;padding:.8rem 1rem;'
                    f'margin:.4rem 0;text-align:right;margin-left:20%">'
                    f'<span style="font-size:.85rem;color:#1a73e8;font-weight:600">You</span><br>'
                    f'<span style="font-size:.9rem">{msg["content"]}</span></div>',
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    f'<div style="background:#f0f4ff;border-radius:12px;padding:.8rem 1rem;'
                    f'margin:.4rem 0;margin-right:20%">'
                    f'<span style="font-size:.85rem;color:#0f3460;font-weight:600">🤖 Claude</span><br>'
                    f'<span style="font-size:.9rem">{msg["content"]}</span></div>',
                    unsafe_allow_html=True)

        # Chat input
        prefill = st.session_state.pop("chat_input_prefill", "")
        user_input = st.text_input(
            "💬 Tell Claude how to improve your resume...",
            value=prefill,
            placeholder='e.g. "Make the summary more aggressive" · "Add cloud computing keywords" · "Shorten experience section"',
            key="chat_input"
        )

        send_col, clear_col = st.columns([3, 1])
        with send_col:
            send = st.button("📨 Send", use_container_width=True, key="chat_send")
        with clear_col:
            if st.button("🗑️ Clear Chat", use_container_width=True, key="chat_clear"):
                st.session_state["chat_history"] = []
                st.session_state.pop("chat_resume", None)
                st.rerun()

        if send and user_input.strip():
            with st.spinner("🤖 Claude is updating your resume..."):
                try:
                    client = anthropic.Anthropic(api_key=api_key)

                    chat_sys = """You are a professional resume editor assistant.
The user will give you instructions to modify their resume.

STRICT RULES:
- Zero markdown: no ##, no **, no *, no --, no hashtags, no asterisks
- Section headers in ALL CAPS only
- Bullets use • only
- Experience: Job Title   Company Name   Start - End  (ONE LINE)
- Education: Degree   University   Start - End  (ONE LINE)
- Always return the COMPLETE updated resume, not just the changed parts
- Never hallucinate new companies, dates, or achievements
- Only modify what the user asks — keep everything else exactly the same

After the resume, on a NEW LINE write:
---CHAT---
Then a brief 1-2 sentence friendly explanation of what you changed."""

                    # Build message history for multi-turn context
                    messages = []
                    for h in st.session_state["chat_history"]:
                        messages.append({"role": h["role"], "content": h["content"]})

                    # Add current user message with resume context
                    full_user_msg = f"""Current Resume:
{current_resume}

User Request: {user_input}

Return the complete updated resume, then ---CHAT--- then explain what changed."""

                    messages.append({"role": "user", "content": full_user_msg})

                    resp = client.messages.create(
                        model="claude-opus-4-5", max_tokens=2500,
                        system=chat_sys,
                        messages=messages
                    )

                    full_response = resp.content[0].text

                    # Split resume from chat explanation
                    if "---CHAT---" in full_response:
                        new_resume, chat_reply = full_response.split("---CHAT---", 1)
                        new_resume  = new_resume.strip()
                        chat_reply  = chat_reply.strip()
                    else:
                        new_resume  = full_response.strip()
                        chat_reply  = "Resume updated as requested."

                    # Save updated resume
                    st.session_state["chat_resume"] = new_resume

                    # Add to chat history (store short versions for display)
                    st.session_state["chat_history"].append({"role": "user",      "content": user_input})
                    st.session_state["chat_history"].append({"role": "assistant", "content": f"✅ Done! {chat_reply}"})

                    st.rerun()

                except Exception as e:
                    st.error(f"Error: {e}")

        # Show updated resume preview + download
        if "chat_resume" in st.session_state:
            st.markdown("---")
            st.markdown("### 📄 Updated Resume Preview")
            st.markdown(render_html(st.session_state["chat_resume"]), unsafe_allow_html=True)

            st.markdown("#### 📋 Copy")
            st.text_area("Ctrl+A → Ctrl+C", value=st.session_state["chat_resume"], height=80, key="chat_copy")

            st.markdown("#### 📥 Download Updated Resume")
            cname  = st.session_state.get("candidate_name", "Resume")
            jtitle = st.session_state.get("job_title", "Role")
            fname3 = f"{cname.replace(' ','_')}_{jtitle.replace(' ','_')}_Chat"
            dc1, dc2 = st.columns(2)
            with dc1:
                st.download_button("📝 Download Word (.docx)",
                    data=build_docx(st.session_state["chat_resume"]),
                    file_name=f"{fname3}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            with dc2:
                st.download_button("📕 Download PDF",
                    data=build_pdf(st.session_state["chat_resume"]),
                    file_name=f"{fname3}.pdf",
                    mime="application/pdf",
                    use_container_width=True)

            if st.button("🔄 Use This as New Base Resume", use_container_width=True, key="chat_set_base"):
                st.session_state["resume_text"] = st.session_state["chat_resume"]
                st.success("✅ Set as your main resume! Go to Tab 2 to re-check ATS score.")

# ═══════════════════════════════════════════
# TAB 4 — Prompt Log
# ═══════════════════════════════════════════
with tab4:
    st.markdown("### 📜 Prompt Engineering Evolution Log")
    st.markdown("Required for IE 5300 — documents prompt evolution from simple to complex.")
    for v in [
        {"v":"v1 — Basic Prompt","s":"⭐⭐",
         "sys":"You are a resume writer. Write a professional resume.",
         "usr":"Write a resume for {name}. Skills: {skills}. Experience: {experience}.",
         "p":"Generic output, no formatting, ignored job requirements entirely."},
        {"v":"v2 — Structured Prompt","s":"⭐⭐⭐",
         "sys":"You are an expert resume writer. Create structured resumes with clear sections: Contact, Summary, Education, Skills, Experience, Projects.",
         "usr":"Name, Email, Education, Skills, Experience, Projects, Target Role.",
         "p":"Better structure but not job-tailored. No keyword matching or ATS optimization."},
        {"v":"v3 — Job-Tailored + ATS + Keyword Injection","s":"⭐⭐⭐⭐⭐",
         "sys":"No markdown. Extract 12-15 keywords from JD. Weave ALL keywords into bullet points. Job Title + Company + Date on ONE line always. Action verbs + metrics. NEVER hallucinate.",
         "usr":"Full candidate profile + complete job description + tone + length preference.",
         "p":"Current best version. Keyword injection ensures high ATS match scores. Future: add few-shot examples."},
    ]:
        with st.expander(f"{v['v']}  {v['s']}"):
            st.markdown("**🤖 System Prompt:**"); st.code(v["sys"], language="text")
            st.markdown("**👤 User Prompt:**");   st.code(v["usr"], language="text")
            st.markdown(f"**⚠️ Issue / What Changed:** {v['p']}")

# ═══════════════════════════════════════════
# TAB 5 — How to Use
# ═══════════════════════════════════════════
with tab5:
    st.markdown("### ℹ️ How to Use This App")
    st.markdown("""
**Step 1** → Paste your Anthropic API key in the sidebar.

**Step 2** → Your info is pre-filled. Update anything needed.

**Step 3** → Find a job on Indeed/LinkedIn → paste full job description on the right panel.

**Step 4** → Click **✨ Generate Tailored Resume**. Takes ~10–15 seconds.

**Step 5** → Review the hallucination check result.

**Step 6** → Download as **Word (.docx)** or **PDF**, or copy the text.

**Step 7** → Go to **Tab 2** → paste job description → click **Check ATS Score** for keyword analysis.

**Step 8** → Go to **Tab 3 (Resume Chat)** → type instructions to refine your resume in real time:
- *"Make the summary more aggressive"*
- *"Add cloud computing keywords"*
- *"Shorten the experience section"*
- Or use the Quick Action buttons for one-click improvements.
    """)
    st.info("💡 **Project Tip:** Try all 3 prompt versions from the sidebar and screenshot the differences — great evidence for your Evaluation grade component!")
